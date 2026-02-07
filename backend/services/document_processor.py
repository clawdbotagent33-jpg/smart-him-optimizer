"""
문서 처리 서비스
다양한 파일 포맷에서 텍스트를 추출하는 기능 제공
"""
import os
import io
from typing import Optional, Dict, Any
from datetime import datetime
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """다양한 문서 포맷 처리기"""

    # 지원하는 파일 형식
    SUPPORTED_TEXT_FORMATS = {'.txt', '.md', '.csv', '.json', '.xml', '.html', '.htm'}
    SUPPORTED_PDF_FORMATS = {'.pdf'}
    SUPPORTED_WORD_FORMATS = {'.docx', '.doc'}
    SUPPORTED_EXCEL_FORMATS = {'.xlsx', '.xls'}
    SUPPORTED_IMAGE_FORMATS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'}

    @classmethod
    def get_supported_extensions(cls) -> set:
        """지원하는 모든 파일 확장자 반환"""
        return (
            cls.SUPPORTED_TEXT_FORMATS |
            cls.SUPPORTED_PDF_FORMATS |
            cls.SUPPORTED_WORD_FORMATS |
            cls.SUPPORTED_EXCEL_FORMATS |
            cls.SUPPORTED_IMAGE_FORMATS
        )

    @classmethod
    def extract_text(cls, file_path: str, filename: str) -> Dict[str, Any]:
        """
        파일에서 텍스트 추출

        Args:
            file_path: 파일 경로
            filename: 원본 파일명

        Returns:
            dict: {
                'text': 추출된 텍스트,
                'metadata': 파일 메타데이터,
                'page_count': 페이지 수 (해당 시),
                'error': 오류 메시지 (있는 경우)
            }
        """
        file_ext = os.path.splitext(filename)[1].lower()

        result = {
            'text': '',
            'metadata': {
                'filename': filename,
                'extension': file_ext,
                'processed_at': datetime.utcnow().isoformat()
            },
            'page_count': 0,
            'error': None
        }

        try:
            if file_ext in cls.SUPPORTED_TEXT_FORMATS:
                result.update(cls._extract_from_text(file_path))
            elif file_ext in cls.SUPPORTED_PDF_FORMATS:
                result.update(cls._extract_from_pdf(file_path))
            elif file_ext in cls.SUPPORTED_WORD_FORMATS:
                result.update(cls._extract_from_word(file_path))
            elif file_ext in cls.SUPPORTED_EXCEL_FORMATS:
                result.update(cls._extract_from_excel(file_path))
            elif file_ext in cls.SUPPORTED_IMAGE_FORMATS:
                result.update(cls._extract_from_image(file_path))
            else:
                result['error'] = f"지원하지 않는 파일 형식: {file_ext}"
                logger.warning(f"Unsupported file format: {file_ext}")

        except Exception as e:
            result['error'] = f"문서 처리 오류: {str(e)}"
            logger.error(f"Error processing document {filename}: {e}", exc_info=True)

        return result

    @classmethod
    def _extract_from_text(cls, file_path: str) -> Dict[str, Any]:
        """텍스트 파일에서 내용 추출"""
        encodings = ['utf-8', 'cp949', 'euc-kr', 'latin1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                return {'text': text, 'page_count': 1}
            except UnicodeDecodeError:
                continue

        # 모든 인코딩 실패 시 기본 처리
        with open(file_path, 'rb') as f:
            content = f.read()
        return {'text': content.decode('utf-8', errors='ignore'), 'page_count': 1}

    @classmethod
    def _extract_from_pdf(cls, file_path: str) -> Dict[str, Any]:
        """PDF에서 텍스트 추출 (PyPDF2)"""
        try:
            import PyPDF2

            text_parts = []
            page_count = 0

            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                page_count = len(pdf_reader.pages)

                for page in pdf_reader.pages:
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(text)
                    except Exception as e:
                        logger.warning(f"Error extracting text from PDF page: {e}")
                        continue

            return {
                'text': '\n\n'.join(text_parts),
                'page_count': page_count
            }
        except ImportError:
            return {
                'text': '',
                'page_count': 0,
                'error': 'PyPDF2 라이브러리가 설치되지 않았습니다'
            }
        except Exception as e:
            return {
                'text': '',
                'page_count': 0,
                'error': f'PDF 처리 오류: {str(e)}'
            }

    @classmethod
    def _extract_from_pdf_with_pdfplumber(cls, file_path: str) -> Dict[str, Any]:
        """pdfplumber를 사용한 PDF 처리 (PyPDF2 실패 시 대안)"""
        try:
            import pdfplumber

            text_parts = []
            page_count = 0

            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)

                for page in pdf.pages:
                    try:
                        text = page.extract_text()
                        if text and text.strip():
                            text_parts.append(text)
                    except Exception:
                        continue

            return {
                'text': '\n\n'.join(text_parts),
                'page_count': page_count
            }
        except ImportError:
            return {
                'text': '',
                'page_count': 0,
                'error': 'pdfplumber 라이브러리가 설치되지 않았습니다'
            }

    @classmethod
    def _extract_from_word(cls, file_path: str) -> Dict[str, Any]:
        """Word 문서에서 텍스트 추출"""
        file_ext = os.path.splitext(file_path)[1].lower()

        # .docx 처리
        if file_ext == '.docx':
            try:
                from docx import Document

                doc = Document(file_path)
                text_parts = []

                # 문단 추출
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)

                # 표 추출
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join([cell.text for cell in row.cells])
                        if row_text.strip():
                            text_parts.append(row_text)

                return {
                    'text': '\n'.join(text_parts),
                    'page_count': 1
                }
            except ImportError:
                return {
                    'text': '',
                    'page_count': 0,
                    'error': 'python-docx 라이브러리가 설치되지 않았습니다'
                }
            except Exception as e:
                return {
                    'text': '',
                    'page_count': 0,
                    'error': f'Word 문서 처리 오류: {str(e)}'
                }

        # .doc 처리 (구버전)
        elif file_ext == '.doc':
            return {
                'text': '',
                'page_count': 0,
                'error': '구버전 Word(.doc) 파일은 지원하지 않습니다. .docx로 변환해주세요.'
            }

    @classmethod
    def _extract_from_excel(cls, file_path: str) -> Dict[str, Any]:
        """Excel 파일에서 텍스트 추출"""
        file_ext = os.path.splitext(file_path)[1].lower()

        try:
            # .xlsx 처리
            if file_ext == '.xlsx':
                import openpyxl

                wb = openpyxl.load_workbook(file_path, data_only=True)
                text_parts = []
                sheet_count = len(wb.sheetnames)

                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    text_parts.append(f"[시트: {sheet_name}]")

                    for row in ws.iter_rows(values_only=True):
                        row_text = ' | '.join([str(cell) if cell is not None else '' for cell in row])
                        if row_text.strip() and row_text != ' | '.join([''] * len(row)):
                            text_parts.append(row_text)

                return {
                    'text': '\n'.join(text_parts),
                    'page_count': sheet_count
                }

            # .xls 처리
            elif file_ext == '.xls':
                import xlrd

                wb = xlrd.open_workbook(file_path)
                text_parts = []
                sheet_count = len(wb.sheet_names())

                for sheet in wb.sheets():
                    text_parts.append(f"[시트: {sheet.name}]")

                    for row_idx in range(sheet.nrows):
                        row = sheet.row(row_idx)
                        row_text = ' | '.join([str(cell.value) if cell.value else '' for cell in row])
                        if row_text.strip():
                            text_parts.append(row_text)

                return {
                    'text': '\n'.join(text_parts),
                    'page_count': sheet_count
                }

        except ImportError as e:
            missing_lib = 'openpyxl' if file_ext == '.xlsx' else 'xlrd'
            return {
                'text': '',
                'page_count': 0,
                'error': f'{missing_lib} 라이브러리가 설치되지 않았습니다'
            }
        except Exception as e:
            return {
                'text': '',
                'page_count': 0,
                'error': f'Excel 처리 오류: {str(e)}'
            }

    @classmethod
    def _extract_from_image(cls, file_path: str) -> Dict[str, Any]:
        """이미지에서 텍스트 추출 (OCR)"""
        try:
            # OCR 라이브러리 확인
            try:
                import pytesseract
                from PIL import Image
            except ImportError:
                return {
                    'text': '',
                    'page_count': 0,
                    'error': 'OCR 기능을 사용하려면 pytesseract와 Pillow가 필요합니다.'
                }

            # 이미지 열기
            image = Image.open(file_path)

            # OCR 수행
            # Korean + English로 설정
            text = pytesseract.image_to_string(image, lang='kor+eng')

            return {
                'text': text,
                'page_count': 1
            }

        except Exception as e:
            return {
                'text': '',
                'page_count': 0,
                'error': f'이미지 OCR 처리 오류: {str(e)}'
            }

    @classmethod
    def sanitize_text(cls, text: str) -> str:
        """
        추출된 텍스트 정제
        - 불필요한 공백 제거
        - 특수 문자 정리
        """
        if not text:
            return ''

        # 연속된 공백 줄이기
        text = ' '.join(text.split())

        # 연속된 개행문 2개로 줄이기
        text = '\n\n'.join([p.strip() for p in text.split('\n') if p.strip()])

        return text.strip()

    @classmethod
    def process_document(cls, file_path: str, filename: str, sanitize: bool = True) -> Dict[str, Any]:
        """
        문서 처리 및 텍스트 추출 (메인 메서드)

        Args:
            file_path: 파일 경로
            filename: 파일명
            sanitize: 텍스트 정제 여부

        Returns:
            처리 결과 dict
        """
        result = cls.extract_text(file_path, filename)

        if result.get('error'):
            return result

        # 텍스트 정제
        if sanitize and result.get('text'):
            result['text'] = cls.sanitize_text(result['text'])

        # 텍스트 길이 로깅
        text_length = len(result.get('text', ''))
        logger.info(f"문서 처리 완료: {filename}, 추출 텍스트 길이: {text_length}")

        return result


# 싱글톤 함수
def extract_text_from_file(file_path: str, filename: str) -> tuple[str, Dict[str, Any]]:
    """
    파일에서 텍스트 추출 (편의 함수)

    Returns:
        (추출된 텍스트, 메타데이터)
    """
    processor = DocumentProcessor()
    result = processor.process_document(file_path, filename)

    text = result.get('text', '')
    metadata = result.get('metadata', {})

    if result.get('error'):
        metadata['error'] = result['error']

    return text, metadata
